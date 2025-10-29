"""
Functions for extracting "blocks" from a .docx file:
 - paragraphs (text + basic formatting)
 - tables (list of rows)
 - images (sha1, size, rel_id) — attempts to preserve links with paragraphs
"""
from pathlib import Path
from typing import Any, Dict, List
from docx import Document
import hashlib
import logging

_LOGGER = logging.getLogger(__name__)


def _safe_hex_color(run) -> str:
    """Returns color in #RRGGBB format or default #000000."""
    try:
        color = run.font.color
        if color is not None and getattr(color, "rgb", None):
            rgb = color.rgb  # RGBColor
            return f"#{rgb[0]:02x}{rgb[1]:02x}{rgb[2]:02x}"
    except Exception:
        pass
    return "#000000"

def extract_docx_blocks(path: Path) -> List[Dict[str, Any]]:
    """
    Returns a list of blocks in the order they appear in the document.
    Each block has a 'type' field (paragraph|table|image) and corresponding data.
    """
    path = Path(path)
    if not path.exists():
        raise FileNotFoundError(f"File does not exist: {path}")

    doc = Document(str(path))
    blocks: List[Dict[str, Any]] = []

    # Mapping of related parts (images and others) for quick access
    related = getattr(doc.part, "related_parts", {})

    # Namespaces for parsing blip elements (drawingML)
    namespaces = {
        "a": "http://schemas.openxmlformats.org/drawingml/2006/main",
        "r": "http://schemas.openxmlformats.org/officeDocument/2006/relationships",
        "wp": "http://schemas.openxmlformats.org/drawingml/2006/wordprocessingDrawing",
        "pic": "http://schemas.openxmlformats.org/drawingml/2006/picture",
    }

    # Iterate through elements in the document body while preserving order
    for element in doc.element.body:
        tag = element.tag
        if tag.endswith("}p"):  # paragraph
            # find the corresponding Paragraph object
            para_obj = None
            for para in doc.paragraphs:
                if para._element is element:
                    para_obj = para
                    break
            if para_obj is None:
                continue
            text = para_obj.text.strip()
            if not text:
                # One could consider adding empty paragraphs (e.g., spacing), but we skip them here
                continue

            # basic info from the first run
            first_run = para_obj.runs[0] if para_obj.runs else None
            blocks.append(
                {
                    "type": "paragraph",
                    "text": text,
                    "style": para_obj.style.name if para_obj.style else "Normal",
                    "bold": bool(first_run and first_run.bold),
                    "italic": bool(first_run and first_run.italic),
                    "underline": bool(first_run and first_run.underline),
                    "color": _safe_hex_color(first_run) if first_run else "#000000",
                }
            )

            # try to find images embedded in runs and add them right after the paragraph
            for run in para_obj.runs:
                # look for blip tags with r:embed reference (r:id)
                try:
                    blips = run._element.findall(".//a:blip", namespaces)
                    for blip in blips:
                        embed = blip.get("{http://schemas.openxmlformats.org/officeDocument/2006/relationships}embed")
                        if embed and embed in related:
                            part = related[embed]
                            data = part.blob if hasattr(part, "blob") else part._blob
                            h = hashlib.sha1(data).hexdigest()
                            blocks.append(
                                {
                                    "type": "image",
                                    "rel_id": embed,
                                    "sha1": h,
                                    "size": len(data),
                                    "filename": getattr(part, "partname", str(embed)),
                                }
                            )
                except Exception:
                    _LOGGER.debug("Error while extracting image from run", exc_info=True)

        elif tag.endswith("}tbl"):  # table
            tbl_obj = None
            for tbl in doc.tables:
                if tbl._element is element:
                    tbl_obj = tbl
                    break
            if tbl_obj is None:
                continue
            rows: List[List[str]] = []
            for row in tbl_obj.rows:
                rows.append([cell.text.strip() for cell in row.cells])
            blocks.append({"type": "table", "table": rows})

    return blocks
